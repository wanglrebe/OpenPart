# backend/app/plugins/crawler_base.py
"""
OpenPart 爬虫插件基础规范

这个文件定义了所有爬虫插件必须遵循的接口标准
"""

from abc import ABC, abstractmethod
from datetime import datetime
from pydantic import BaseModel, Field
from enum import Enum
import tempfile
import os
import time
from typing import Optional, Dict, List, Any, Union
from pathlib import Path

class PluginStatus(Enum):
    """插件状态枚举"""
    ACTIVE = "active"           # 活跃
    INACTIVE = "inactive"       # 未激活
    ERROR = "error"            # 错误状态
    DISABLED = "disabled"       # 已禁用

class DataSourceType(Enum):
    """数据源类型枚举"""
    ECOMMERCE = "ecommerce"     # 电商平台
    SUPPLIER = "supplier"       # 供应商网站
    DATABASE = "database"       # 数据库
    API = "api"                # API接口
    CATALOG = "catalog"        # 产品目录
    DOCUMENT = "document"      # 文档文件（新增）
    OTHER = "other"            # 其他

class PartData(BaseModel):
    """爬取的零件数据标准格式"""
    name: str = Field(..., description="零件名称，必填")
    category: Optional[str] = Field(None, description="零件类别")
    description: Optional[str] = Field(None, description="零件描述")
    properties: Optional[Dict[str, Any]] = Field(None, description="自定义属性")
    image_url: Optional[str] = Field(None, description="图片URL")
    source_url: Optional[str] = Field(None, description="原始数据URL")
    external_id: Optional[str] = Field(None, description="外部系统ID")
    price: Optional[float] = Field(None, description="价格信息")
    availability: Optional[str] = Field(None, description="库存状态")
    
    class Config:
        json_encoders = {
            datetime: lambda v: v.isoformat()
        }

class ConfigField(BaseModel):
    """配置字段定义"""
    name: str = Field(..., description="字段名称")
    label: str = Field(..., description="显示标签")
    type: str = Field(..., description="字段类型: text|password|number|select|checkbox|textarea")
    required: bool = Field(False, description="是否必填")
    default: Any = Field(None, description="默认值")
    placeholder: Optional[str] = Field(None, description="占位符文本")
    options: Optional[List[Dict[str, str]]] = Field(None, description="选择项(select类型使用)")
    help_text: Optional[str] = Field(None, description="帮助说明")
    validation: Optional[Dict[str, Any]] = Field(None, description="验证规则")

class PluginInfo(BaseModel):
    """插件信息"""
    name: str = Field(..., description="插件名称")
    version: str = Field(..., description="版本号")
    description: str = Field(..., description="插件描述")
    author: str = Field(..., description="作者")
    data_source: str = Field(..., description="数据源名称")
    data_source_type: DataSourceType = Field(..., description="数据源类型")
    homepage: Optional[str] = Field(None, description="数据源主页")
    terms_url: Optional[str] = Field(None, description="服务条款URL")
    rate_limit: Optional[int] = Field(None, description="请求频率限制(秒)")
    batch_size: Optional[int] = Field(50, description="批次处理大小")

class CrawlResult(BaseModel):
    """爬取结果"""
    success: bool = Field(..., description="是否成功")
    data: List[PartData] = Field(default_factory=list, description="爬取的数据")
    total_count: int = Field(0, description="总数据量")
    error_message: Optional[str] = Field(None, description="错误信息")
    warnings: List[str] = Field(default_factory=list, description="警告信息")
    execution_time: Optional[float] = Field(None, description="执行时间(秒)")
    next_page_token: Optional[str] = Field(None, description="下一页标识")

class TestResult(BaseModel):
    """测试结果"""
    success: bool = Field(..., description="测试是否成功")
    message: str = Field(..., description="测试结果信息")
    response_time: Optional[float] = Field(None, description="响应时间(秒)")
    sample_data: Optional[Dict[str, Any]] = Field(None, description="示例数据")

class BaseCrawlerPlugin(ABC):
    """
    爬虫插件基类
    
    所有爬虫插件都必须继承此类并实现所有抽象方法
    """
    
    def __init__(self):
        """初始化插件"""
        self._validate_plugin_info()
        self._admin_token = None  # 添加token属性
        self._api_base = "http://localhost:8000/api"  # 添加API基础地址
    
    # 添加token相关方法
    def set_admin_token(self, token: str):
        """设置管理员token"""
        self._admin_token = token
    
    def get_admin_token(self) -> str:
        """获取管理员token"""
        return self._admin_token
    
    def get_api_base(self) -> str:
        """获取API基础地址"""
        return self._api_base
    
    def get_auth_headers(self) -> Dict[str, str]:
        """获取认证请求头"""
        if self._admin_token:
            return {"Authorization": f"Bearer {self._admin_token}"}
        return {}
    
    # 添加文件下载辅助方法
    # 在 crawler_base.py 中，修改 download_uploaded_file 方法，添加调试信息：

    def download_uploaded_file(self, file_id: str) -> bytes:
        """下载上传的文件 - 添加调试信息"""
        import requests
        
        if not self._admin_token:
            raise ValueError("缺少管理员token，无法下载文件")
        
        print(f"准备下载文件: {file_id}")
        print(f"使用token (前50字符): {self._admin_token[:50]}...")
        print(f"API地址: {self._api_base}/admin/files/{file_id}/download")
        
        try:
            headers = self.get_auth_headers()
            print(f"请求头: {headers}")
            
            response = requests.get(
                f"{self._api_base}/admin/files/{file_id}/download",
                headers=headers,
                timeout=30
            )
            
            print(f"响应状态码: {response.status_code}")
            print(f"响应头: {dict(response.headers)}")
            
            if response.status_code == 200:
                print(f"文件下载成功，大小: {len(response.content)} bytes")
                return response.content
            else:
                print(f"文件下载失败，响应内容: {response.text[:200]}...")
                raise Exception(f"文件下载失败: HTTP {response.status_code} - {response.text}")
                
        except Exception as e:
            print(f"下载文件时出现异常: {str(e)}")
            raise Exception(f"下载文件时出错: {str(e)}")
    
    @property
    @abstractmethod
    def plugin_info(self) -> PluginInfo:
        """
        返回插件基本信息
        
        Returns:
            PluginInfo: 插件信息对象
        """
        pass
    
    @property
    @abstractmethod
    def config_schema(self) -> List[ConfigField]:
        """
        返回插件配置表单定义
        
        Returns:
            List[ConfigField]: 配置字段列表
        """
        pass
    
    @abstractmethod
    def validate_config(self, config: Dict[str, Any]) -> bool:
        """
        验证配置参数是否有效
        
        Args:
            config: 配置参数字典
            
        Returns:
            bool: 配置是否有效
            
        Raises:
            ValueError: 配置无效时抛出异常，异常信息会显示给用户
        """
        pass
    
    @abstractmethod
    def test_connection(self, config: Dict[str, Any]) -> TestResult:
        """
        测试与数据源的连接
        
        Args:
            config: 配置参数字典
            
        Returns:
            TestResult: 测试结果
        """
        pass
    
    @abstractmethod
    def crawl(self, config: Dict[str, Any], **kwargs) -> CrawlResult:
        """
        执行数据爬取
        
        Args:
            config: 配置参数字典
            **kwargs: 额外参数，可能包含:
                - page_token: 分页标识
                - limit: 数据限制数量
                - filters: 过滤条件
                
        Returns:
            CrawlResult: 爬取结果
        """
        pass
    
    def get_allowed_domains(self) -> List[str]:
        """
        返回插件允许访问的域名列表
        
        Returns:
            List[str]: 允许访问的域名列表
        """
        return []
    
    def get_required_permissions(self) -> List[str]:
        """
        返回插件需要的权限列表
        
        Returns:
            List[str]: 权限列表，可选值:
                - network: 网络访问
                - file_read: 文件读取
                - file_write: 文件写入
        """
        return ["network"]
    
    def cleanup(self):
        """
        插件卸载时的清理工作
        
        在插件被卸载或禁用时调用，用于释放资源
        """
        pass
    
    def _validate_plugin_info(self):
        """验证插件信息"""
        try:
            info = self.plugin_info
            if not info.name or not info.version:
                raise ValueError("插件名称和版本不能为空")
        except Exception as e:
            raise ValueError(f"插件信息验证失败: {e}")

# 安全装饰器
def safe_network_call(allowed_domains: List[str] = None):
    """
    网络调用安全装饰器
    
    Args:
        allowed_domains: 允许访问的域名列表
    """
    def decorator(func):
        def wrapper(*args, **kwargs):
            # 这里会在实际执行时添加网络访问检查
            return func(*args, **kwargs)
        return wrapper
    return decorator

# 工具类
class PluginUtils:
    """插件开发工具类 - 增强版本"""
    
    # 文件大小限制
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    MAX_PAGES_PDF = 100  # PDF最大页数
    MAX_ROWS_EXCEL = 10000  # Excel最大行数
    TEMP_FILE_LIFETIME = 3600  # 临时文件生存时间（秒）
    
    @staticmethod
    def clean_text(text: str) -> str:
        """清理文本内容"""
        if not text:
            return ""
        return text.strip().replace('\n', ' ').replace('\r', '')
    
    @staticmethod
    def extract_number(text: str) -> Optional[float]:
        """从文本中提取数字"""
        import re
        if not text:
            return None
        
        numbers = re.findall(r'\d+\.?\d*', str(text))
        if numbers:
            try:
                return float(numbers[0])
            except ValueError:
                return None
        return None
    
    @staticmethod
    def normalize_url(url: str, base_url: str = None) -> str:
        """标准化URL"""
        from urllib.parse import urljoin, urlparse
        
        if not url:
            return ""
        
        if base_url and not urlparse(url).netloc:
            return urljoin(base_url, url)
        
        return url
    
    @staticmethod
    def validate_url(url: str) -> bool:
        """验证URL格式"""
        from urllib.parse import urlparse
        
        try:
            result = urlparse(url)
            return all([result.scheme, result.netloc])
        except Exception:
            return False
        
    # ==================== 新增：安全文件处理方法 ====================
    
    @staticmethod
    def _create_temp_file(content: bytes, suffix: str = '') -> str:
        """创建临时文件"""
        if len(content) > PluginUtils.MAX_FILE_SIZE:
            raise PluginError(f"文件大小超过限制 ({PluginUtils.MAX_FILE_SIZE // 1024 // 1024}MB)")
        
        # 创建临时文件
        temp_fd, temp_path = tempfile.mkstemp(suffix=suffix, prefix='plugin_')
        try:
            with os.fdopen(temp_fd, 'wb') as f:
                f.write(content)
            return temp_path
        except Exception as e:
            # 清理失败的临时文件
            try:
                os.unlink(temp_path)
            except:
                pass
            raise PluginError(f"创建临时文件失败: {str(e)}")
    
    @staticmethod
    def _cleanup_temp_file(file_path: str):
        """清理临时文件"""
        try:
            if file_path and os.path.exists(file_path):
                os.unlink(file_path)
        except Exception as e:
            print(f"清理临时文件失败: {e}")
    
    @staticmethod
    def safe_read_pdf(file_content: bytes, **options) -> str:
        """
        安全的PDF读取
        
        Args:
            file_content: PDF文件二进制内容
            **options: 
                - max_pages: 最大页数 (默认100)
                - extract_tables: 是否提取表格 (默认False)
                - password: PDF密码 (如果需要)
        
        Returns:
            str: 提取的文本内容
        """
        temp_path = None
        try:
            # 创建临时文件
            temp_path = PluginUtils._create_temp_file(file_content, '.pdf')
            
            # 导入PDF库
            try:
                import pdfplumber
            except ImportError:
                raise PluginError("pdfplumber 库未安装，请联系管理员")
            
            max_pages = min(options.get('max_pages', PluginUtils.MAX_PAGES_PDF), PluginUtils.MAX_PAGES_PDF)
            extract_tables = options.get('extract_tables', False)
            password = options.get('password')
            
            content_parts = []
            
            with pdfplumber.open(temp_path, password=password) as pdf:
                total_pages = len(pdf.pages)
                pages_to_process = min(total_pages, max_pages)
                
                print(f"处理PDF: 总页数 {total_pages}, 处理页数 {pages_to_process}")
                
                for i, page in enumerate(pdf.pages[:pages_to_process]):
                    # 提取文本
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"=== 第 {i+1} 页 ===\n{text}")
                    
                    # 提取表格（如果需要）
                    if extract_tables:
                        tables = page.extract_tables()
                        for j, table in enumerate(tables):
                            if table:
                                table_text = f"\n--- 表格 {j+1} ---\n"
                                for row in table:
                                    if row:
                                        table_text += " | ".join(str(cell or '') for cell in row) + "\n"
                                content_parts.append(table_text)
            
            result = "\n\n".join(content_parts)
            print(f"PDF解析完成，提取文本长度: {len(result)} 字符")
            return result
            
        except Exception as e:
            raise PluginError(f"PDF读取失败: {str(e)}")
        finally:
            if temp_path:
                PluginUtils._cleanup_temp_file(temp_path)
    
    @staticmethod
    def safe_read_excel(file_content: bytes, **options) -> List[Dict[str, Any]]:
        """
        安全的Excel读取
        
        Args:
            file_content: Excel文件二进制内容
            **options:
                - sheet_name: 工作表名称 (默认第一个)
                - max_rows: 最大行数 (默认10000)
                - header_row: 表头行号 (默认0)
        
        Returns:
            List[Dict]: 行数据列表
        """
        temp_path = None
        try:
            # 创建临时文件
            temp_path = PluginUtils._create_temp_file(file_content, '.xlsx')
            
            try:
                import pandas as pd
            except ImportError:
                raise PluginError("pandas 库未安装，请联系管理员")
            
            max_rows = min(options.get('max_rows', PluginUtils.MAX_ROWS_EXCEL), PluginUtils.MAX_ROWS_EXCEL)
            sheet_name = options.get('sheet_name', 0)  # 默认第一个工作表
            header_row = options.get('header_row', 0)
            
            # 读取Excel
            df = pd.read_excel(
                temp_path, 
                sheet_name=sheet_name,
                nrows=max_rows,
                header=header_row
            )
            
            # 清理数据
            df = df.fillna('')  # 填充空值
            
            print(f"Excel解析完成，读取 {len(df)} 行数据")
            
            return df.to_dict('records')
            
        except Exception as e:
            raise PluginError(f"Excel读取失败: {str(e)}")
        finally:
            if temp_path:
                PluginUtils._cleanup_temp_file(temp_path)
    
    @staticmethod
    def safe_read_word(file_content: bytes, **options) -> str:
        """
        安全的Word文档读取
        
        Args:
            file_content: Word文件二进制内容
            **options:
                - extract_tables: 是否提取表格 (默认True)
                - extract_images: 是否提取图片信息 (默认False)
        
        Returns:
            str: 提取的文本内容
        """
        temp_path = None
        try:
            temp_path = PluginUtils._create_temp_file(file_content, '.docx')
            
            try:
                from docx import Document
            except ImportError:
                raise PluginError("python-docx 库未安装，请联系管理员")
            
            extract_tables = options.get('extract_tables', True)
            
            doc = Document(temp_path)
            content_parts = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            # 提取表格（如果需要）
            if extract_tables:
                for i, table in enumerate(doc.tables):
                    table_text = f"\n=== 表格 {i+1} ===\n"
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            table_text += row_text + "\n"
                    content_parts.append(table_text)
            
            result = "\n".join(content_parts)
            print(f"Word文档解析完成，提取文本长度: {len(result)} 字符")
            return result
            
        except Exception as e:
            raise PluginError(f"Word文档读取失败: {str(e)}")
        finally:
            if temp_path:
                PluginUtils._cleanup_temp_file(temp_path)
    
    @staticmethod
    def safe_read_csv(file_content: bytes, **options) -> List[Dict[str, Any]]:
        """
        安全的CSV读取
        
        Args:
            file_content: CSV文件内容
            **options:
                - encoding: 文件编码 (默认自动检测)
                - delimiter: 分隔符 (默认自动检测)
                - max_rows: 最大行数
        
        Returns:
            List[Dict]: 行数据列表
        """
        try:
            import pandas as pd
        except ImportError:
            raise PluginError("pandas 库未安装，请联系管理员")
        
        try:
            # 尝试不同编码
            encodings = [options.get('encoding'), 'utf-8', 'gbk', 'gb2312', 'utf-8-sig']
            text_content = None
            
            for encoding in encodings:
                if encoding is None:
                    continue
                try:
                    text_content = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                text_content = file_content.decode('utf-8', errors='ignore')
            
            # 检测分隔符
            delimiter = options.get('delimiter')
            if not delimiter:
                # 简单的分隔符检测
                line = text_content.split('\n')[0]
                if '\t' in line:
                    delimiter = '\t'
                elif ',' in line:
                    delimiter = ','
                elif ';' in line:
                    delimiter = ';'
                else:
                    delimiter = ','
            
            max_rows = min(options.get('max_rows', PluginUtils.MAX_ROWS_EXCEL), PluginUtils.MAX_ROWS_EXCEL)
            
            # 使用StringIO读取CSV
            from io import StringIO
            df = pd.read_csv(
                StringIO(text_content),
                delimiter=delimiter,
                nrows=max_rows
            )
            
            # 清理数据
            df = df.fillna('')
            
            print(f"CSV解析完成，读取 {len(df)} 行数据")
            return df.to_dict('records')
            
        except Exception as e:
            raise PluginError(f"CSV读取失败: {str(e)}")
    
    @staticmethod
    def detect_file_type(file_content: bytes) -> str:
        """
        检测文件类型
        
        Returns:
            str: 文件类型 ('pdf', 'excel', 'word', 'csv', 'unknown')
        """
        if not file_content:
            return 'unknown'
        
        # 检查文件头
        if file_content.startswith(b'%PDF'):
            return 'pdf'
        elif file_content.startswith(b'PK\x03\x04'):
            # ZIP格式，可能是Office文档
            if b'word/' in file_content[:1000]:
                return 'word'
            elif b'xl/' in file_content[:1000]:
                return 'excel'
            else:
                return 'zip'
        elif file_content.startswith(b'\xd0\xcf\x11\xe0'):
            # 老版本Office格式
            return 'office_legacy'
        else:
            # 尝试作为文本解析
            try:
                text = file_content[:1000].decode('utf-8', errors='ignore')
                if ',' in text or '\t' in text:
                    return 'csv'
                else:
                    return 'text'
            except:
                return 'unknown'
    
    @staticmethod
    def safe_parse_any_file(file_content: bytes, **options) -> Dict[str, Any]:
        """
        智能解析任意类型的文档文件
        
        Args:
            file_content: 文件二进制内容
            **options: 解析选项
        
        Returns:
            Dict: 包含解析结果和元信息
        """
        try:
            file_type = PluginUtils.detect_file_type(file_content)
            
            result = {
                'file_type': file_type,
                'file_size': len(file_content),
                'content': None,
                'success': False,
                'error': None
            }
            
            try:
                if file_type == 'pdf':
                    result['content'] = PluginUtils.safe_read_pdf(file_content, **options)
                    result['success'] = True
                    
                elif file_type == 'excel':
                    result['content'] = PluginUtils.safe_read_excel(file_content, **options)
                    result['success'] = True
                    
                elif file_type == 'word':
                    result['content'] = PluginUtils.safe_read_word(file_content, **options)
                    result['success'] = True
                    
                elif file_type in ['csv', 'text']:
                    result['content'] = PluginUtils.safe_read_csv(file_content, **options)
                    result['success'] = True
                    
                else:
                    result['error'] = f"不支持的文件类型: {file_type}"
                    
            except Exception as e:
                result['error'] = str(e)
            
            return result
            
        except Exception as e:
            return {
                'file_type': 'unknown',
                'file_size': len(file_content),
                'content': None,
                'success': False,
                'error': f"文件解析失败: {str(e)}"
            }

# 异常类定义
class PluginError(Exception):
    """插件基础异常"""
    pass

class ConfigError(PluginError):
    """配置错误"""
    pass

class NetworkError(PluginError):
    """网络错误"""
    pass

class DataError(PluginError):
    """数据错误"""
    pass

class SecurityError(PluginError):
    """安全错误"""
    pass